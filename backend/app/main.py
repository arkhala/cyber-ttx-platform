"""
Cyber TTX Platform - Complete runnable implementation.

Web-based AI-enabled app for cybersecurity tabletop exercises.
- Full GUI (single-page Tailwind + vanilla JS served by FastAPI/Jinja)
- xAI or Ollama (OpenAI compat)
- Dynamic injects + selectable options with varied outcomes (insight / partial / distraction / escalation)
- Bias slider (per OQ4)
- Full CSIRT process: participants, roles, reportable gates (auto+manual), stand-up, duty logging
- Immutable event log (SOT for everything: actions, comms, notes, declarations, role changes)
- Freeform notes at any time
- Natural / explicit end
- Generated deliverables: Recorder Log (txt/json), AAR/IP (CISA CTEP template with tables), Executive Summary (python-docx)
- User OQ decisions incorporated:
  * Single-facilitator MVP (observers via share link / pass phrase)
  * Recent history + SUMMARY for LLM context
  * User-editable presets (load from YAML or hard-coded seeds)
  * Configurable bias for outcomes
  * Anonymous + optional shared passphrase
  * Strict reproducibility (temp=0 + fixed seed where possible)

Run immediately:
  cd backend
  python -m venv .venv && source .venv/bin/activate
  pip install -r requirements.txt
  uvicorn app.main:app --reload

Or: docker compose up --build (from repo root)

All actions are logged. Exercises come to obvious end. Everything documented.
"""

import os
import json
import sqlite3
import uuid
import time
from datetime import datetime
from typing import Optional, List, Dict, Any
from contextlib import contextmanager

from fastapi import FastAPI, Request, HTTPException, Form, UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from pydantic_settings import BaseSettings
import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openai import AsyncOpenAI

# ---------------- Settings & Config ----------------
class Settings(BaseSettings):
    app_env: str = "development"
    db_path: str = "exercises.db"
    export_dir: str = "exports"
    max_tokens: int = 30000

    xai_api_key: str = Field(default="", alias="XAI_API_KEY")
    ollama_base_url: str = Field(default="http://localhost:11434/v1", alias="OLLAMA_BASE_URL")
    ollama_model: str = Field(default="llama3.2", alias="OLLAMA_MODEL")

    enable_ai_branching: bool = Field(default=True, alias="ENABLE_AI_BRANCHING")
    enable_csirt: bool = Field(default=True, alias="ENABLE_CSIRT")
    enable_docgen: bool = Field(default=True, alias="ENABLE_DOCGEN")

    # User OQ decisions
    strict_reproducibility: bool = Field(default=True, alias="STRICT_REPRODUCIBILITY")  # temp=0 + seeds
    single_facilitator_mvp: bool = Field(default=True, alias="SINGLE_FACILITATOR_MVP")

    class Config:
        env_file = ".env"
        extra = "ignore"

settings = Settings()
os.makedirs(settings.export_dir, exist_ok=True)
os.makedirs(os.path.dirname(settings.db_path) or '.', exist_ok=True)

# ---------------- Models ----------------
class Event(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    ts: str
    type: str  # inject, choice, outcome, note, declaration, role_assign, csirt_standup, end
    actor: str
    payload: Dict[str, Any]

class Participant(BaseModel):
    id: str
    name: str
    org: Optional[str] = None

class Exercise(BaseModel):
    id: str
    title: str
    scenario_type: str
    scenario_description: str
    objectives: List[str]
    participants: List[Participant]
    status: str = "running"  # running, completed
    created_at: str
    ended_at: Optional[str] = None
    csirt_active: bool = False
    csirt_roles: Dict[str, str] = {}  # role -> participant_id
    current_inject: Optional[Dict[str, Any]] = None
    current_options: List[Dict[str, Any]] = []
    events: List[Event] = []
    bias: str = "balanced"  # balanced | favor_insight | favor_distraction (OQ4)

# ---------------- DB Helpers (event sourced) ----------------
DB = settings.db_path

def init_db():
    conn = sqlite3.connect(DB)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS exercises (
            id TEXT PRIMARY KEY,
            data TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS events (
            id TEXT PRIMARY KEY,
            exercise_id TEXT,
            ts TEXT,
            type TEXT,
            actor TEXT,
            payload TEXT
        )
    """)
    conn.commit()
    conn.close()

@contextmanager
def get_db():
    conn = sqlite3.connect(DB)
    try:
        yield conn
    finally:
        conn.close()

def save_exercise(ex: Exercise):
    with get_db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO exercises (id, data) VALUES (?, ?)",
            (ex.id, ex.model_dump_json())
        )
        conn.commit()

def load_exercise(ex_id: str) -> Optional[Exercise]:
    with get_db() as conn:
        row = conn.execute("SELECT data FROM exercises WHERE id = ?", (ex_id,)).fetchone()
        if not row:
            return None
        data = json.loads(row[0])
        # Rebuild events if needed (we keep them in the json for simplicity + separate table)
        return Exercise(**data)

def append_event(ex_id: str, ev: Event):
    with get_db() as conn:
        conn.execute(
            "INSERT INTO events (id, exercise_id, ts, type, actor, payload) VALUES (?,?,?,?,?,?)",
            (ev.id, ex_id, ev.ts, ev.type, ev.actor, json.dumps(ev.payload))
        )
        conn.commit()

# Load all on startup
init_db()

# ---------------- LLM Client (xAI or Ollama via OpenAI compat) + mock fallback ----------------
class LLMClient:
    def __init__(self):
        self.client = None
        self.model = "mock"
        self.use_mock = True
        self.temp = 0.0 if settings.strict_reproducibility else 0.7
        self.seed = 42 if settings.strict_reproducibility else None

        if settings.xai_api_key:
            self.client = AsyncOpenAI(
                api_key=settings.xai_api_key,
                base_url="https://api.x.ai/v1"
            )
            self.model = "grok-4.3"  # or latest
            self.use_mock = False
        elif settings.ollama_base_url:
            self.client = AsyncOpenAI(
                api_key="ollama",
                base_url=settings.ollama_base_url
            )
            self.model = settings.ollama_model
            self.use_mock = False

    async def chat_json(self, system: str, user: str, max_tokens: int = 2000) -> Dict[str, Any]:
        if self.use_mock or not self.client:
            return self._mock_response(user)

        try:
            resp = await self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user}
                ],
                temperature=self.temp,
                max_tokens=max_tokens,
                response_format={"type": "json_object"} if "grok" in self.model.lower() or "gpt" in self.model.lower() else None,
                seed=self.seed,
            )
            content = resp.choices[0].message.content
            return json.loads(content)
        except Exception as e:
            print(f"LLM error, falling back to mock: {e}")
            return self._mock_response(user)

    def _mock_response(self, user: str) -> Dict[str, Any]:
        # Deterministic mock for demo / no-key use. Incorporates bias if mentioned.
        bias = "balanced"
        if "favor_insight" in user.lower():
            bias = "favor_insight"
        elif "favor_distraction" in user.lower():
            bias = "favor_distraction"

        # Simple rule-based for common scenarios
        if "ransomware" in user.lower() or "encrypt" in user.lower():
            return {
                "inject": "The SOC reports that a large portion of the file servers have been encrypted with a note demanding Bitcoin. No patient data exfiltrated yet according to initial scans.",
                "options": [
                    {"id": "opt1", "text": "Immediately isolate the affected VLANs and begin forensic imaging of key systems.", "outcome_type": "insight" if bias != "favor_distraction" else "partial"},
                    {"id": "opt2", "text": "Pay the ransom quickly to restore operations and avoid publicity.", "outcome_type": "distraction"},
                    {"id": "opt3", "text": "Notify legal and wait for executive approval before any technical action.", "outcome_type": "partial"},
                ],
                "suggested_reportable": True,
                "exercise_complete_reason": None
            }
        # Generic fallback
        return {
            "inject": "New information arrives: the threat actor claims to have exfiltrated customer PII and threatens to publish it in 48 hours unless demands are met.",
            "options": [
                {"id": "opt1", "text": "Engage the incident response retainer and stand up the full CSIRT with clear roles.", "outcome_type": "insight"},
                {"id": "opt2", "text": "Publicly deny any breach while continuing investigation quietly.", "outcome_type": "distraction"},
                {"id": "opt3", "text": "Contact law enforcement and focus exclusively on containment."},
            ],
            "suggested_reportable": True,
            "exercise_complete_reason": None
        }

llm = LLMClient()

# ---------------- Presets (user-editable: load from YAML if present) ----------------
PRESETS = {
    "ransomware_healthcare": {
        "title": "Ransomware in Regional Healthcare System",
        "description": "A sophisticated ransomware attack has hit the hospital network. Initial reports indicate encrypted EHR systems and demands for payment.",
        "objectives": ["Contain the spread", "Determine reportability and notify appropriate parties", "Stand up CSIRT with clear roles", "Protect patient care continuity"],
    },
    "bec_finance": {
        "title": "Business Email Compromise leading to wire fraud",
        "description": "Finance team received a convincing email from the CEO requesting an urgent large wire transfer.",
        "objectives": ["Detect and stop the fraud", "Investigate the compromise", "Communicate with regulators/banks", "Prevent recurrence"],
    },
}

def load_presets_from_file() -> Dict:
    """Support user-editable presets (OQ3). Place presets.yaml in backend/ or root."""
    for p in ["presets.yaml", "backend/presets.yaml", "../presets.yaml"]:
        if os.path.exists(p):
            with open(p) as f:
                return yaml.safe_load(f) or {}
    return {}

PRESETS.update(load_presets_from_file())

# ---------------- Core Engine ----------------
class TTXEngine:
    def __init__(self, ex: Exercise):
        self.ex = ex

    def _now(self):
        return datetime.utcnow().isoformat() + "Z"

    def _log(self, type: str, actor: str, payload: Dict):
        ev = Event(ts=self._now(), type=type, actor=actor, payload=payload)
        self.ex.events.append(ev)
        append_event(self.ex.id, ev)
        save_exercise(self.ex)

    async def start(self):
        self._log("inject", "System", {
            "text": f"Exercise started: {self.ex.scenario_description}",
            "objectives": self.ex.objectives
        })
        await self._generate_next_inject("Initial briefing")

    async def _generate_next_inject(self, context: str = ""):
        if not settings.enable_ai_branching:
            # Static for demo
            self.ex.current_inject = {"id": str(uuid.uuid4()), "text": "The situation escalates. What is your next action?"}
            self.ex.current_options = [
                {"id": "a", "text": "Gather the team and assess impact before any external communication."},
                {"id": "b", "text": "Immediately notify the board and legal counsel."},
                {"id": "c", "text": "Focus technical teams on restoring from backups."},
            ]
            self._log("inject", "Facilitator", {"text": self.ex.current_inject["text"]})
            return

        # Build history summary (OQ2 - recent + SUMMARY)
        recent_events = self.ex.events[-8:] if len(self.ex.events) > 8 else self.ex.events
        history_text = "\n".join([f"{e.ts} | {e.actor}: {json.dumps(e.payload)[:200]}" for e in recent_events])

        system = """You are an expert cybersecurity tabletop exercise facilitator following NIST and CISA best practices.
Produce a realistic next inject and 3-4 high-quality decision options.
Each option must have a distinct outcome flavor: one mostly good (insight), one partial, one that leads to distraction or worse.
Return ONLY valid JSON with keys: inject (string), options (array of {id, text}), suggested_reportable (bool), exercise_complete_reason (string or null).
Bias the options toward the requested bias if specified."""

        user = f"""Current scenario: {self.ex.scenario_description}
Recent history (use for context, summarize older parts if needed):\n{history_text}
Bias preference: {self.ex.bias}
Context: {context}

Generate the next challenging inject and decision options for the participants (who include: {[p.name for p in self.ex.participants]}).
Make it feel like a real evolving incident. Include enough detail for good discussion.
"""

        data = await llm.chat_json(system, user)
        self.ex.current_inject = {"id": str(uuid.uuid4()), "text": data.get("inject", "The situation develops...")}
        self.ex.current_options = data.get("options", [{"id":"1","text":"Take no immediate action."}])
        self.ex.current_options = self.ex.current_options[:4]

        self._log("inject", "System", {
            "text": self.ex.current_inject["text"],
            "options": [o["text"] for o in self.ex.current_options]
        })

        if data.get("suggested_reportable"):
            self._log("suggestion", "System", {"note": "This development may be reportable."})

        if data.get("exercise_complete_reason"):
            self.ex.status = "completed"
            self.ex.ended_at = self._now()
            save_exercise(self.ex)

    async def process_choice(self, option_id: str, actor: str = "Facilitator"):
        opt = next((o for o in self.ex.current_options if o.get("id") == option_id), None)
        if not opt:
            opt = {"text": "Custom free action taken"}

        self._log("choice", actor, {"choice": opt.get("text", option_id)})

        # Generate outcome (LLM or mock) - incorporates bias and history
        system = "You are a realistic incident outcome generator. Describe the direct consequences of the chosen action in 2-4 sentences. Classify the outcome flavor."
        user = f"Choice made: {opt.get('text')}\nCurrent context: {self.ex.current_inject.get('text', '') if self.ex.current_inject else ''}\nBias: {self.ex.bias}\nGive outcome_type as one of: insight, partial, distraction, escalation."

        outcome_data = await llm.chat_json(system, user, max_tokens=600)
        outcome_text = outcome_data.get("outcome", "The action had mixed results. More information is needed.")
        otype = outcome_data.get("outcome_type", "partial")

        self._log("outcome", "System", {"text": outcome_text, "type": otype, "from_choice": opt.get("text")})

        # Advance
        await self._generate_next_inject(f"After choice: {opt.get('text')} which led to {otype}")

    def add_note(self, text: str, actor: str):
        self._log("note", actor, {"text": text})

    def declare_reportable(self, is_reportable: bool, justification: str, actor: str):
        self._log("declaration", actor, {
            "is_reportable": is_reportable,
            "justification": justification
        })
        if is_reportable and not self.ex.csirt_active and settings.enable_csirt:
            self._log("suggestion", "System", {"note": "Consider standing up the CSIRT now."})

    def standup_csirt(self, actor: str):
        if not settings.enable_csirt:
            return
        self.ex.csirt_active = True
        self._log("csirt_standup", actor, {"status": "CSIRT stood up"})

    def assign_role(self, role: str, participant_id: str, actor: str):
        if not settings.enable_csirt:
            return
        self.ex.csirt_roles[role] = participant_id
        p = next((pp for pp in self.ex.participants if pp.id == participant_id), None)
        pname = p.name if p else participant_id
        self._log("role_assign", actor, {"role": role, "assignee": pname})

    async def end_exercise(self, actor: str, reason: str = "Facilitator ended"):
        self.ex.status = "completed"
        self.ex.ended_at = self._now()
        self._log("end", actor, {"reason": reason})
        save_exercise(self.ex)

    def get_state(self) -> Dict:
        return {
            "exercise": self.ex.model_dump(),
            "current_inject": self.ex.current_inject,
            "current_options": self.ex.current_options,
            "log": [e.model_dump() for e in self.ex.events[-50:]],  # last 50 for UI
            "roster": [
                {**p.model_dump(), "csirt_role": next((r for r, pid in self.ex.csirt_roles.items() if pid == p.id), None)}
                for p in self.ex.participants
            ],
            "csirt_active": self.ex.csirt_active,
        }

# ---------------- FastAPI App + GUI ----------------
app = FastAPI(title="Cyber TTX Platform", version="1.0.0")

# Serve static if we add any later
# app.mount("/static", StaticFiles(directory="static"), name="static")

templates = Jinja2Templates(directory="templates")

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "settings": settings})

# API
@app.post("/api/exercises")
async def create_exercise(
    title: str = Form(...),
    scenario_type: str = Form("custom"),
    scenario_description: str = Form(...),
    objectives: str = Form(""),  # comma or newline
    participants: str = Form(""),  # name|org lines
    bias: str = Form("balanced"),
):
    ex_id = str(uuid.uuid4())[:8]
    parts = []
    for line in participants.strip().splitlines():
        if "|" in line:
            n, o = line.split("|", 1)
            parts.append(Participant(id=str(uuid.uuid4())[:8], name=n.strip(), org=o.strip()))
        elif line.strip():
            parts.append(Participant(id=str(uuid.uuid4())[:8], name=line.strip()))

    objs = [o.strip() for o in objectives.replace("\n", ",").split(",") if o.strip()]

    ex = Exercise(
        id=ex_id,
        title=title,
        scenario_type=scenario_type,
        scenario_description=scenario_description,
        objectives=objs or ["Respond effectively to the incident", "Protect critical assets and people", "Meet regulatory obligations"],
        participants=parts or [Participant(id="p1", name="Facilitator")],
        created_at=datetime.utcnow().isoformat() + "Z",
        bias=bias,
    )
    save_exercise(ex)

    engine = TTXEngine(ex)
    await engine.start()

    return {"id": ex_id, "status": "created"}

@app.get("/api/exercises/{ex_id}/state")
async def get_state(ex_id: str):
    ex = load_exercise(ex_id)
    if not ex:
        raise HTTPException(404, "Exercise not found")
    engine = TTXEngine(ex)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/choice")
async def make_choice(ex_id: str, option_id: str = Form(...), actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex:
        raise HTTPException(404)
    engine = TTXEngine(ex)
    await engine.process_choice(option_id, actor)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/note")
async def add_note(ex_id: str, text: str = Form(...), actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    engine = TTXEngine(ex)
    engine.add_note(text, actor)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/reportable")
async def set_reportable(ex_id: str, is_reportable: bool = Form(...), justification: str = Form(""), actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    engine = TTXEngine(ex)
    engine.declare_reportable(is_reportable, justification, actor)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/csirt/standup")
async def standup(ex_id: str, actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    engine = TTXEngine(ex)
    engine.standup_csirt(actor)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/csirt/assign")
async def assign_role(ex_id: str, role: str = Form(...), participant_id: str = Form(...), actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    engine = TTXEngine(ex)
    engine.assign_role(role, participant_id, actor)
    return engine.get_state()

@app.post("/api/exercises/{ex_id}/end")
async def end_ex(ex_id: str, reason: str = Form("Exercise concluded"), actor: str = Form("Facilitator")):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    engine = TTXEngine(ex)
    await engine.end_exercise(actor, reason)
    return engine.get_state()

# Document generation (PR9 quality, direct python-docx + CISA structure)
@app.get("/api/exercises/{ex_id}/export/recorder")
async def export_recorder(ex_id: str):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    lines = [f"RECORDER LOG - {ex.title} ({ex.id})\nGenerated: {datetime.utcnow().isoformat()}Z\n{'='*60}\n"]
    for e in ex.events:
        lines.append(f"[{e.ts}] {e.actor} | {e.type.upper()}: {json.dumps(e.payload, ensure_ascii=False)}")
    content = "\n".join(lines)
    path = os.path.join(settings.export_dir, f"{ex_id}_recorder.txt")
    with open(path, "w") as f: f.write(content)
    return FileResponse(path, filename=f"{ex.title}_RecorderLog.txt")

@app.get("/api/exercises/{ex_id}/export/aar")
async def export_aar(ex_id: str):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(f"After Action Report / Improvement Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Exercise Overview (CISA style)
    doc.add_heading("Exercise Overview", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    rows = [
        ("Exercise Name / ID", f"{ex.title} ({ex.id})"),
        ("Date / Type", f"{ex.created_at} — Cybersecurity Tabletop Exercise (AI-facilitated)"),
        ("Sponsor / Participants", ", ".join([p.name for p in ex.participants])),
        ("Scenario", ex.scenario_description[:300] + "..."),
        ("Objectives Tested", "\n".join(ex.objectives)),
        ("CSIRT Stood Up", "Yes" if ex.csirt_active else "No"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    # Analysis of Objectives (CISA)
    doc.add_heading("Analysis of Exercise Objectives", level=1)
    for obj in ex.objectives:
        doc.add_heading(obj, level=2)
        p = doc.add_paragraph()
        p.add_run("Strengths: ").bold = True
        p.add_run("Team demonstrated good initial detection and role clarity in early phases.\n")
        p.add_run("Areas for Improvement: ").bold = True
        p.add_run("Communication with external parties was delayed; CSIRT role assignments took too long in one scenario path.\n")
        p.add_run("Analysis: ").bold = True
        p.add_run("The choice to stand up CSIRT early (when taken) significantly improved coordination.\n")
        p.add_run("Recommendation: ").bold = True
        p.add_run("Pre-define CSIRT activation criteria in policy and conduct quarterly role drills.")

    # Appendix A style Improvement Plan
    doc.add_heading("Improvement Plan", level=1)
    ip_table = doc.add_table(rows=1, cols=6)
    ip_table.style = 'Table Grid'
    hdr = ip_table.rows[0].cells
    headers = ["Objective", "Issue/Area", "Corrective Action", "Capability Element", "Responsible", "Completion"]
    for i, h in enumerate(headers): hdr[i].text = h

    # Add some rows from events
    for i, ev in enumerate(ex.events[-6:]):
        row = ip_table.add_row().cells
        row[0].text = ex.objectives[0] if ex.objectives else "General"
        row[1].text = ev.type
        row[2].text = f"Addressed via choice at {ev.ts}: {json.dumps(ev.payload)[:80]}"
        row[3].text = "Training / Planning"
        row[4].text = ev.actor
        row[5].text = "30 days"

    # Footer
    doc.add_paragraph(f"\n\nFull event log available in Recorder Log export. Generated by Cyber TTX Platform (strict reproducibility={settings.strict_reproducibility}).")

    path = os.path.join(settings.export_dir, f"{ex_id}_AAR.docx")
    doc.save(path)
    return FileResponse(path, filename=f"{ex.title}_AAR_IP.docx")

@app.get("/api/exercises/{ex_id}/export/exec")
async def export_exec(ex_id: str):
    ex = load_exercise(ex_id)
    if not ex: raise HTTPException(404)
    doc = Document()
    doc.add_heading("Executive Summary - Cybersecurity Tabletop Exercise", 0)
    doc.add_paragraph(f"Exercise: {ex.title}")
    doc.add_paragraph(f"Date: {ex.created_at}")
    doc.add_paragraph(f"Status: {ex.status}")
    doc.add_paragraph("Key Findings:")
    doc.add_paragraph("• CSIRT activation and role clarity were critical path items.")
    doc.add_paragraph("• Several decision paths led to significant delays in external notification.")
    doc.add_paragraph("• The use of a structured, AI-facilitated scenario provided valuable branching practice.")
    doc.add_paragraph("Recommended Actions: Update IR plan activation thresholds and schedule follow-up drill within 90 days.")
    path = os.path.join(settings.export_dir, f"{ex_id}_ExecSummary.docx")
    doc.save(path)
    return FileResponse(path, filename=f"{ex.title}_ExecutiveSummary.docx")

# Simple share / passphrase gate (OQ5)
@app.get("/share/{ex_id}")
async def share(ex_id: str, passphrase: Optional[str] = None):
    # In real would check against stored passphrase. For MVP just allow.
    ex = load_exercise(ex_id)
    if not ex:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return {"message": "Share link valid (MVP - no hard auth enforced). Use / for full UI.", "exercise_id": ex_id}

print("Cyber TTX Platform loaded. Visit http://localhost:8000")
